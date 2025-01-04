from fastapi import FastAPI
from fastapi.responses import FileResponse, JSONResponse
from typing import List
from docx import Document
from pydantic import BaseModel
import os

app = FastAPI()

# Пути к шаблонам
TEMPLATES = {
    "contract": "templates/бланк_Массив.docx",
    "schet": "templates/бланк_Счет.docx",  # Новый шаблон для счета
}

# Описание структуры записи
class Record(BaseModel):
    name_yl: str = ""
    role_s: str = ""
    fio_s: str = ""
    adress_yl: str = ""
    role: str = ""
    fio: str = ""
    price: str = ""
    inn: str = ""

# Функция для замены маркеров в тексте
def replace_markers_in_runs(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    for marker, value in replacements.items():
        if marker in full_text:
            full_text = full_text.replace(marker, value)

    if full_text:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full_text

# Функция заполнения шаблона
def fill_template(template_path: str, record: Record, output_path: str):
    # Открываем документ по указанному пути
    document = Document(template_path)

    replacements = {
        "{{name_yl}}": record.name_yl,
        "{{role_s}}": record.role_s,
        "{{fio_s}}": record.fio_s,
        "{{adress_yl}}": record.adress_yl,
        "{{role}}": record.role,
        "{{fio}}": record.fio,
        "{{price}}": record.price,
        "{{inn}}": record.inn,
    }

    # Замена маркеров в параграфах
    for paragraph in document.paragraphs:
        replace_markers_in_runs(paragraph, replacements)

    # Замена маркеров в таблицах
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_markers_in_runs(paragraph, replacements)

    # Сохраняем заполненный документ
    document.save(output_path)

# Общая функция для генерации файлов
def generate_documents(template_key: str, records: List[Record]):
    try:
        # Проверяем, что указанный шаблон существует
        if template_key not in TEMPLATES:
            return {"error": f"Template '{template_key}' not found"}

        template_path = TEMPLATES[template_key]
        output_files = []

        # Создаем папку для выходных файлов
        os.makedirs("output", exist_ok=True)

        # Генерация файлов для каждого набора данных
        for idx, record in enumerate(records, start=1):
            output_filename = f"output/{template_key}_сформированный{idx}.docx"
            fill_template(template_path, record, output_filename)
            output_files.append(output_filename)

        return output_files

    except Exception as e:
        return {"error": str(e)}

# Эндпоинт для договоров
@app.post("/generate/")
async def generate_contract(records: List[Record]):
    files = generate_documents("contract", records)
    if "error" in files:
        return JSONResponse(files)

    # Если один файл — возвращаем его напрямую
    if len(files) == 1:
        return FileResponse(
            files[0],
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(files[0]),
        )

    # Если файлов несколько — возвращаем их список
    return JSONResponse({"files": files})

# Эндпоинт для счетов
@app.post("/schet/")
async def generate_schet(records: List[Record]):
    files = generate_documents("schet", records)
    if "error" in files:
        return JSONResponse(files)

    # Если один файл — возвращаем его напрямую
    if len(files) == 1:
        return FileResponse(
            files[0],
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(files[0]),
        )

    # Если файлов несколько — возвращаем их список
    return JSONResponse({"files": files})
